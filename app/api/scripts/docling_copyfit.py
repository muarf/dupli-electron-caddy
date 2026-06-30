import sys
import docx
from docx.shared import Pt, Inches
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
import math

def generate_docx(pdf_path, docx_path):
    print(f"Starting copyfitting for {pdf_path}")
    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False
    converter = DocumentConverter(format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)})
    res = converter.convert(pdf_path)
    doc = res.document
    
    document = docx.Document()
    
    # Try to extract page dimensions from docling doc
    pdf_width = 8.27 * 72 # default to A4 in points
    pdf_height = 11.69 * 72
    
    if hasattr(doc, 'pages') and len(doc.pages) > 0:
        first_page = list(doc.pages.values())[0]
        if hasattr(first_page, 'size'):
            pdf_width = getattr(first_page.size, 'width', pdf_width)
            pdf_height = getattr(first_page.size, 'height', pdf_height)
            
    # Set page size
    section = document.sections[0]
    # Convert points (72 pt = 1 inch) to Inches for docx
    section.page_width = Inches(pdf_width / 72.0)
    section.page_height = Inches(pdf_height / 72.0)
    
    # Set small margins to allow absolute-like positioning via indents
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    current_page = 1
    
    for item, level in doc.iterate_items():
        text = getattr(item, 'text', '')
        if not text:
            continue
            
        import re
        text = re.sub(r' +', ' ', text)
            
        page_no = 1
        bbox = None
        if hasattr(item, 'prov') and len(item.prov) > 0:
            page_no = item.prov[0].page_no
            bbox = item.prov[0].bbox
            
        while current_page < page_no:
            document.add_page_break()
            current_page += 1
            
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        estimated_pt = 11.0
        label = getattr(item, 'label', '')
        if label:
            label_str = str(label).lower()
            if 'title' in label_str:
                estimated_pt = 18.0
            elif 'section' in label_str or 'header' in label_str:
                estimated_pt = 14.0
            elif 'caption' in label_str or 'footnote' in label_str or 'page_header' in label_str or 'page_footer' in label_str:
                estimated_pt = 9.0
        
        if bbox:
            try:
                left = getattr(bbox, 'l', 0)
                # Assume Docling coordinates are in points (72 per inch)
                # But Docling origin might be bottom-left. The width/height works regardless.
                if left > 36: # If more than 0.5 inches from left edge
                    indent_inches = (left - 36) / 72.0
                    if indent_inches > 0:
                        p.paragraph_format.left_indent = Inches(indent_inches)
            except Exception as e:
                print(f"Bbox parsing error: {e}")
                
        run = p.add_run(text)
        run.font.size = Pt(estimated_pt)
        # TODO: Detect original font, fallback to Arial for now
        run.font.name = 'Arial'
        
        # Line spacing adjusted to font size
        p.paragraph_format.line_spacing = Pt(estimated_pt * 1.15)

    document.save(docx_path)
    print(f"Saved DOCX to {docx_path}")

if __name__ == '__main__':
    generate_docx(sys.argv[1], sys.argv[2])
